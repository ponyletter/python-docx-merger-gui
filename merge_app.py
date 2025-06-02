import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docxcompose.composer import Composer

# --- Configuration ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_INPUT_DIR = os.path.join(SCRIPT_DIR, "input_docs")
DEFAULT_OUTPUT_DIR = os.path.join(SCRIPT_DIR, "output_docs")
DEFAULT_OUTPUT_FILENAME = "merged_document.docx"

# Ensure default directories exist
os.makedirs(DEFAULT_INPUT_DIR, exist_ok=True)
os.makedirs(DEFAULT_OUTPUT_DIR, exist_ok=True)


def merge_docx_files(source_files_list, target_file_path):
    """
    Merges a list of .docx files into a single target .docx file.

    Args:
        source_files_list (list): A list of full paths to .docx files to merge.
                                  The order in this list determines merge order.
        target_file_path (str): The full path for the merged output .docx file.

    Returns:
        str: A message indicating success or failure.
    """
    if not source_files_list:
        return "错误：没有找到要合并的 .docx 文件。"
    if len(source_files_list) < 1: # Technically could merge 1 file (copy it)
        return "错误：至少需要一个文件才能执行操作。"

    try:
        # Use the first document as the base for the composer
        master_doc = Document(source_files_list[0])
        composer = Composer(master_doc)

        # If there are more documents, append them with page breaks
        if len(source_files_list) > 1:
            page_break_doc = Document()  # Create a temporary doc for the page break
            page_break_doc.add_page_break()

            for i in range(1, len(source_files_list)):
                composer.append(page_break_doc)  # Add page break
                doc_to_append = Document(source_files_list[i])
                composer.append(doc_to_append)

        composer.save(target_file_path)
        return f"成功：{len(source_files_list)} 个文档已合并到\n{target_file_path}"
    except Exception as e:
        return f"合并过程中发生错误：\n{str(e)}"


class DocxMergerApp:
    def __init__(self, master):
        self.master = master
        master.title("Word 文档合并工具")
        master.geometry("550x350") # Adjusted size for better layout

        # --- Style ---
        style = ttk.Style()
        style.configure("TButton", padding=6, relief="flat", font=('Helvetica', 10))
        style.configure("TLabel", padding=5, font=('Helvetica', 10))
        style.configure("TEntry", padding=5, font=('Helvetica', 10))

        # --- Variables ---
        self.input_dir_var = tk.StringVar(value=DEFAULT_INPUT_DIR)
        self.output_dir_var = tk.StringVar(value=DEFAULT_OUTPUT_DIR)
        self.output_filename_var = tk.StringVar(value=DEFAULT_OUTPUT_FILENAME)
        self.status_var = tk.StringVar(value="请选择文件夹并开始合并。")

        # --- UI Elements ---
        # Input Directory
        ttk.Label(master, text="输入文件夹 (包含 .docx):").grid(row=0, column=0, padx=10, pady=5, sticky="w")
        self.input_entry = ttk.Entry(master, textvariable=self.input_dir_var, width=40)
        self.input_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        ttk.Button(master, text="浏览...", command=self.select_input_dir).grid(row=0, column=2, padx=5, pady=5)

        # Output Directory
        ttk.Label(master, text="输出文件夹:").grid(row=1, column=0, padx=10, pady=5, sticky="w")
        self.output_entry = ttk.Entry(master, textvariable=self.output_dir_var, width=40)
        self.output_entry.grid(row=1, column=1, padx=5, pady=5, sticky="ew")
        ttk.Button(master, text="浏览...", command=self.select_output_dir).grid(row=1, column=2, padx=5, pady=5)

        # Output Filename
        ttk.Label(master, text="输出文件名:").grid(row=2, column=0, padx=10, pady=5, sticky="w")
        self.filename_entry = ttk.Entry(master, textvariable=self.output_filename_var, width=40)
        self.filename_entry.grid(row=2, column=1, padx=5, pady=5, sticky="ew")
        ttk.Label(master, text=".docx").grid(row=2, column=2, padx=0, pady=5, sticky="w")


        # Merge Button
        self.merge_button = ttk.Button(master, text="开始合并", command=self.start_merge_process)
        self.merge_button.grid(row=3, column=0, columnspan=3, padx=10, pady=20)

        # Status Label
        self.status_label = ttk.Label(master, textvariable=self.status_var, wraplength=500, justify="center")
        self.status_label.grid(row=4, column=0, columnspan=3, padx=10, pady=10)

        # Configure column weights for responsiveness
        master.grid_columnconfigure(1, weight=1)

    def select_input_dir(self):
        dir_path = filedialog.askdirectory(initialdir=self.input_dir_var.get(), title="选择输入文件夹")
        if dir_path:
            self.input_dir_var.set(dir_path)

    def select_output_dir(self):
        dir_path = filedialog.askdirectory(initialdir=self.output_dir_var.get(), title="选择输出文件夹")
        if dir_path:
            self.output_dir_var.set(dir_path)

    def start_merge_process(self):
        input_path = self.input_dir_var.get()
        output_path = self.output_dir_var.get()
        output_filename = self.output_filename_var.get()

        if not output_filename.lower().endswith(".docx"):
            output_filename += ".docx"

        if not os.path.isdir(input_path):
            messagebox.showerror("错误", f"输入文件夹不存在：\n{input_path}")
            return

        if not os.path.isdir(output_path):
            try:
                os.makedirs(output_path, exist_ok=True)
            except OSError as e:
                messagebox.showerror("错误", f"无法创建输出文件夹：\n{output_path}\n{e}")
                return

        source_files = []
        try:
            # List files and sort them to ensure consistent order
            # os.listdir's order can be OS-dependent, sorting makes it predictable
            filenames = sorted(os.listdir(input_path))
            for filename in filenames:
                if filename.lower().endswith(".docx"):
                    source_files.append(os.path.join(input_path, filename))
        except Exception as e:
            messagebox.showerror("错误", f"读取输入文件夹时出错：\n{e}")
            return

        if not source_files:
            messagebox.showinfo("提示", "在输入文件夹中没有找到 .docx 文件。")
            self.status_var.set("未找到 .docx 文件。")
            return

        target_file = os.path.join(output_path, output_filename)

        self.status_var.set("正在合并文档...")
        self.master.update_idletasks()  # Update GUI before long task

        result_message = merge_docx_files(source_files, target_file)

        if "成功" in result_message:
            messagebox.showinfo("合并完成", result_message)
        else:
            messagebox.showerror("合并失败", result_message)
        self.status_var.set(result_message.replace("\n", " ")) # Display multi-line msg in status


if __name__ == '__main__':
    root = tk.Tk()
    app = DocxMergerApp(root)
    root.mainloop()